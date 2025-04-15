# Streamlit App: Outlook Email Filter and Export

import streamlit as st
import win32com.client
import pythoncom
from datetime import datetime, timedelta
import os
from docx import Document
from fpdf import FPDF

EXPORT_PATH = os.path.join(os.getcwd(), 'EmailExports')
os.makedirs(EXPORT_PATH, exist_ok=True)

pythoncom.CoInitialize()

outlook = win32com.client.Dispatch("Outlook.Application")
namespace = outlook.GetNamespace("MAPI")
inbox = namespace.GetDefaultFolder(6)

st.title("Outlook Email Explorer & Exporter")

# Filters
start_date = st.date_input("Start Date", datetime.now().date() - timedelta(days=1))
end_date = st.date_input("End Date", datetime.now().date())
sender_filter = st.text_input("Filter by Sender Email")
recipient_filter = st.text_input("Filter by Recipient Email")
subject_filter = st.text_input("Filter by Subject Keyword")

# Fetch emails
messages = inbox.Items
messages.Sort("[ReceivedTime]", True)

filtered_emails = []

for message in messages:
    if message.Class == 43:  # Mail Item
        received_date = message.ReceivedTime.date()

        if received_date < start_date or received_date > end_date:
            continue
        if sender_filter and sender_filter.lower() not in str(message.SenderEmailAddress).lower():
            continue
        if recipient_filter and recipient_filter.lower() not in str(message.To).lower():
            continue
        if subject_filter and subject_filter.lower() not in str(message.Subject).lower():
            continue

        attachments = [att.FileName for att in message.Attachments]

        filtered_emails.append({
            'entry_id': message.EntryID,
            'subject': message.Subject,
            'sender': f"{message.SenderName} <{message.SenderEmailAddress}>",
            'received': message.ReceivedTime,
            'body': message.Body,
            'attachments': attachments
        })

selected_ids = []

if filtered_emails:
    st.write(f"Found {len(filtered_emails)} emails matching your filters")

    for email in filtered_emails:
        checkbox = st.checkbox(f"{email['received']} - {email['subject']} ({email['sender']}) - Attachments: {', '.join(email['attachments']) if email['attachments'] else 'None'}", key=email['entry_id'])
        if checkbox:
            selected_ids.append(email)
else:
    st.write("No emails found for the selected filters.")

# Export selected emails
export_format = st.selectbox("Select Export Format", ["TXT", "Word", "PDF"])

if st.button("Export Selected Emails"):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if export_format == "TXT":
        export_file = os.path.join(EXPORT_PATH, f"Selected_Emails_{timestamp}.txt")
        with open(export_file, 'w', encoding='utf-8') as f:
            for email in selected_ids:
                f.write("=" * 50 + "\n")
                f.write(f"Subject: {email['subject']}\n")
                f.write(f"From: {email['sender']}\n")
                f.write(f"Received: {email['received']}\n")
                f.write(f"Attachments: {', '.join(email['attachments']) if email['attachments'] else 'None'}\n\n")
                f.write(email['body'])
                f.write("\n\n")

    elif export_format == "Word":
        export_file = os.path.join(EXPORT_PATH, f"Selected_Emails_{timestamp}.docx")
        doc = Document()
        for email in selected_ids:
            doc.add_heading(email['subject'], level=1)
            doc.add_paragraph(f"From: {email['sender']}")
            doc.add_paragraph(f"Received: {email['received']}")
            doc.add_paragraph(f"Attachments: {', '.join(email['attachments']) if email['attachments'] else 'None'}")
            doc.add_paragraph(email['body'])
            doc.add_paragraph("=" * 50)
        doc.save(export_file)

    elif export_format == "PDF":
        export_file = os.path.join(EXPORT_PATH, f"Selected_Emails_{timestamp}.pdf")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for email in selected_ids:
            pdf.multi_cell(0, 10, f"Subject: {email['subject']}")
            pdf.multi_cell(0, 10, f"From: {email['sender']}")
            pdf.multi_cell(0, 10, f"Received: {email['received']}")
            pdf.multi_cell(0, 10, f"Attachments: {', '.join(email['attachments']) if email['attachments'] else 'None'}")
            pdf.multi_cell(0, 10, email['body'])
            pdf.multi_cell(0, 10, "=" * 50)
        pdf.output(export_file)

    st.success(f"Export completed successfully! File saved at: {export_file}")
